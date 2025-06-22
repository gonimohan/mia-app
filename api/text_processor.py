import logging
from pathlib import Path
import pandas as pd
import pdfplumber
import docx # python-docx
import openpyxl # For pandas to read xlsx

logger = logging.getLogger(__name__)

def count_words(text: str) -> int:
    """Counts words in a given string."""
    if not text:
        return 0
    return len(text.split())

def extract_text_from_pdf(file_path: str) -> tuple[Optional[str], int]:
    """Extracts text from a PDF file using pdfplumber."""
    text_content = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
        full_text = "\n".join(text_content)
        return full_text, count_words(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {e}")
        return None, 0

def extract_text_from_docx(file_path: str) -> tuple[Optional[str], int]:
    """Extracts text from a DOCX file."""
    try:
        doc = docx.Document(file_path)
        text_content = [para.text for para in doc.paragraphs]
        full_text = "\n".join(text_content)
        return full_text, count_words(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {e}")
        return None, 0

def extract_text_from_txt(file_path: str) -> tuple[Optional[str], int]:
    """Extracts text from a TXT file."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            full_text = f.read()
        return full_text, count_words(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from TXT {file_path}: {e}")
        return None, 0

def extract_text_from_csv(file_path: str) -> tuple[Optional[str], int]:
    """Extracts text from a CSV file by concatenating all string columns."""
    try:
        df = pd.read_csv(file_path, dtype=str, keep_default_na=False) # Read all as string initially
        text_content = []
        for col in df.columns:
            # Only join non-empty string cells
            col_text = df[col].astype(str).str.strip().replace('', pd.NA).dropna().str.cat(sep=' ')
            if col_text:
                 text_content.append(col_text)
        full_text = "\n".join(text_content)
        return full_text, count_words(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from CSV {file_path}: {e}")
        return None, 0

def extract_text_from_excel(file_path: str) -> tuple[Optional[str], int]:
    """Extracts text from an XLSX file by concatenating all string cells from all sheets."""
    try:
        xls = pd.ExcelFile(file_path)
        all_sheets_text = []
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name, dtype=str, keep_default_na=False)
            sheet_text_content = []
            for col in df.columns:
                col_text = df[col].astype(str).str.strip().replace('', pd.NA).dropna().str.cat(sep=' ')
                if col_text:
                    sheet_text_content.append(col_text)
            if sheet_text_content:
                all_sheets_text.append(f"Sheet: {sheet_name}\n" + "\n".join(sheet_text_content))

        full_text = "\n\n".join(all_sheets_text)
        return full_text, count_words(full_text)
    except Exception as e:
        logger.error(f"Error extracting text from Excel {file_path}: {e}")
        return None, 0

def extract_text_from_file(file_path_str: str, file_extension: str, mime_type: Optional[str] = None) -> Optional[Dict[str, Any]]:
    """
    Main function to extract text from a file based on its extension or mime type.
    Returns a dictionary with 'text' and 'word_count' or None if extraction fails.
    """
    file_path = Path(file_path_str)
    if not file_path.exists():
        logger.error(f"File not found for extraction: {file_path_str}")
        return None

    extracted_text = None
    word_c = 0

    logger.info(f"Attempting to extract text from {file_path_str} with extension {file_extension}")

    if file_extension == ".pdf":
        extracted_text, word_c = extract_text_from_pdf(file_path_str)
    elif file_extension == ".docx":
        extracted_text, word_c = extract_text_from_docx(file_path_str)
    elif file_extension == ".txt":
        extracted_text, word_c = extract_text_from_txt(file_path_str)
    elif file_extension == ".csv":
        extracted_text, word_c = extract_text_from_csv(file_path_str)
    elif file_extension == ".xlsx":
        extracted_text, word_c = extract_text_from_excel(file_path_str)
    else:
        logger.warning(f"Unsupported file extension for text extraction: {file_extension} for file {file_path_str}")
        return None

    if extracted_text is not None:
        logger.info(f"Successfully extracted {word_c} words from {file_path_str}.")
        return {"text": extracted_text, "word_count": word_c}
    else:
        logger.error(f"Failed to extract text from {file_path_str}.")
        return None

# Example Usage (for testing this file directly)
if __name__ == '__main__':
    # Create dummy files for testing
    test_dir = Path("_test_files")
    test_dir.mkdir(exist_ok=True)

    # TXT
    txt_file = test_dir / "sample.txt"
    with open(txt_file, "w") as f:
        f.write("This is a sample text file.\nIt has two lines.")

    # PDF (requires a real PDF, pdfplumber cannot create one easily)
    # For now, skip direct PDF creation, assume one exists or test manually

    # DOCX
    try:
        doc = docx.Document()
        doc.add_paragraph("Hello, this is a DOCX file.")
        doc.add_paragraph("Another paragraph here.")
        docx_file = test_dir / "sample.docx"
        doc.save(docx_file)
    except Exception as e:
        print(f"Could not create sample.docx: {e}. python-docx might not be installed or other issues.")
        docx_file = None


    # CSV
    csv_file = test_dir / "sample.csv"
    csv_data = {'col1': ['data1', 'data2'], 'col2': ['infoA', 'infoB']}
    pd.DataFrame(csv_data).to_csv(csv_file, index=False)

    # XLSX
    xlsx_file = test_dir / "sample.xlsx"
    excel_data = {'Sheet1': pd.DataFrame({'A': [1, 2], 'B': ['x', 'y']}),
                  'Sheet2': pd.DataFrame({'C': [3, 4], 'D': ['z', 'w']})}
    with pd.ExcelWriter(xlsx_file) as writer:
        for sheet_name, df_sheet in excel_data.items():
            df_sheet.to_excel(writer, sheet_name=sheet_name, index=False)

    test_files = {
        ".txt": str(txt_file),
        # ".pdf": "path_to_your_sample.pdf", # Replace with an actual PDF path to test
        ".docx": str(docx_file) if docx_file else None,
        ".csv": str(csv_file),
        ".xlsx": str(xlsx_file),
        ".unsupported": "dummy.xyz"
    }

    for ext, f_path in test_files.items():
        if f_path is None and ext == ".docx":
            print(f"Skipping DOCX test as file could not be created.")
            continue
        if ext == ".unsupported": # Create a dummy unsupported file
             with open(test_dir / f_path, "w") as f: f.write("dummy")
             f_path = str(test_dir / f_path)


        print(f"\nTesting extraction for {ext} file: {f_path}")
        if not Path(f_path).exists() and ext != ".unsupported":
            print(f"Test file {f_path} not found, skipping.")
            continue

        result = extract_text_from_file(f_path, ext)
        if result:
            print(f"Extracted: {result['word_count']} words. Preview: '{result['text'][:100]}...'")
        else:
            print("Extraction failed or not supported.")

    # Clean up dummy files
    import shutil
    # shutil.rmtree(test_dir) # Comment out if you want to inspect files
    print(f"\nTest files are in {test_dir}. Remove it manually if needed.")


# --- Text Analysis ---
BUSINESS_KEYWORDS = [
    "market trends",
    "competitive landscape",
    "customer behavior",
    "growth strategy"
]

def analyze_text_keywords(text_content: str) -> Dict[str, bool]:
    """
    Analyzes the text content for the presence of predefined business keywords.
    Performs a case-insensitive search.
    """
    analysis_results = {}
    if not text_content:
        for keyword in BUSINESS_KEYWORDS:
            analysis_results[keyword] = False
        return analysis_results

    lower_text_content = text_content.lower()
    for keyword in BUSINESS_KEYWORDS:
        analysis_results[keyword] = keyword.lower() in lower_text_content

    logger.info(f"Keyword analysis complete. Results: {analysis_results}")
    return analysis_results


if __name__ == '__main__':
    # ... (rest of the existing if __name__ == '__main__' block) ...
    # Add tests for keyword analysis
    sample_text_for_analysis = """
    This document discusses market trends and the competitive landscape.
    Understanding customer behavior is key for our new growth strategy.
    We need to be agile.
    """
    print("\nTesting Keyword Analysis:")
    analysis_res = analyze_text_keywords(sample_text_for_analysis)
    print(f"Analysis on sample text: {analysis_res}")

    analysis_empty = analyze_text_keywords("")
    print(f"Analysis on empty text: {analysis_empty}")

    analysis_none = analyze_text_keywords("No relevant keywords here.")
    print(f"Analysis on text with no keywords: {analysis_none}")
